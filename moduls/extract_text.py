import fitz
import docx
import subprocess
import os
import tempfile
import logging
from pathlib import Path
from typing import Optional, List, Tuple

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Вспомогательные функции ---
def _convert_table_to_markdown(table: docx.table.Table) -> str:
    """Вспомогательная функция для конвертации таблицы docx в Markdown"""
    md_table = ""
    # Заголовки
    header_cells = [cell.text.strip().replace("|", "\\|") for cell in table.rows[0].cells]
    md_table += "| " + " | ".join(header_cells) + " |\n"
    # Разделитель
    md_table += "| " + " | ".join(["---"] * len(header_cells)) + " |\n"
    # Строки данных
    for row in table.rows[1:]:
        row_cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        # Пропускаем пустые строки в таблице (если все ячейки пустые)
        if any(row_cells):
            md_table += "| " + " | ".join(row_cells) + " |\n"
    return md_table

def _is_block_inside_bbox(block_bbox: fitz.Rect, table_bboxes: List[fitz.Rect]) -> bool:
    """
    Проверяет, находится ли текстовый блок (по его координатам)
    внутри какой-либо из границ таблиц.
    """
    for table_bbox in table_bboxes:
        # Простая проверка: содержит ли прямоугольник таблицы данный блок.
        # Можно использовать более сложные проверки на пересечение (intersects)
        # или процент перекрытия, если требуется.
        if table_bbox.contains(block_bbox):
            return True
    return False

def _convert_extracted_table_to_markdown(data: list[list]) -> str:
    """
    Конвертирует данные таблицы (список списков) в строку Markdown.
    """
    if not data:
        return ""

    # Очистка данных: замена None на пустую строку, экранирование "|"
    cleaned_data = []
    num_cols = 0
    for row_idx, row in enumerate(data):
        # PyMuPDF может возвращать None для некоторых строк или ячеек
        if row is None: continue

        current_row = []
        # Иногда table.extract() может вернуть список не-списков для заголовка, обработаем это
        if not isinstance(row, (list, tuple)):
             row = [row] # Оборачиваем в список, если это не список

        for cell in row:
            # Заменяем переносы строк внутри ячейки на пробел или <br> для Markdown
            cell_text = str(cell).replace('\r', '').replace('\n', ' ').replace("|", "\\|").strip() if cell is not None else ""
            current_row.append(cell_text)

        cleaned_data.append(current_row)
        num_cols = max(num_cols, len(current_row))

    if not cleaned_data or num_cols == 0:
        return ""

    # Убедимся, что все строки имеют одинаковое количество столбцов
    # Это важно для валидного Markdown
    for row in cleaned_data:
        row.extend([""] * (num_cols - len(row)))

    md_table_lines = []
    # Заголовок (предполагаем, что первая строка - это заголовок)
    header = cleaned_data[0]
    md_table_lines.append("| " + " | ".join(header) + " |")
    # Разделитель
    md_table_lines.append("| " + " | ".join(["---"] * num_cols) + " |")
    # Тело таблицы
    for row in cleaned_data[1:]:
        md_table_lines.append("| " + " | ".join(row) + " |")

    # Возвращаем таблицу с пустой строкой после для разделения
    return "\n".join(md_table_lines) + "\n"

# --- Обработка PDF ---

def extract_markdown_from_pdf_with_tables(pdf_path: str) -> str:
    """
    Извлекает текст из PDF в Markdown, используя page.find_tables()
    для явной обработки таблиц и page.get_text("blocks") для остального текста.
    """
    final_markdown_content = ""
    try:
        doc = fitz.open(pdf_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_elements = [] # Список для хранения текстовых блоков и Markdown таблиц

            # 1. Найти таблицы и их границы
            # Настройки find_tables можно тюнинговать (strategy, vertical_strategy и т.д.)
            # https://pymupdf.readthedocs.io/en/latest/page.html#Page.find_tables
            tabs = page.find_tables(snap_tolerance=3, join_tolerance=3)
            table_bboxes = [fitz.Rect(t.bbox) for t in tabs.tables]
            logging.info(f"Page {page_num + 1}: Found {len(tabs.tables)} table(s).")

            # 2. Предварительно отрендерить таблицы в Markdown
            rendered_tables = {} # Словарь {индекс_таблицы: строка_markdown}
            table_insertion_points = {} # Словарь {индекс_таблицы: верхняя_координата_y}
            for i, table in enumerate(tabs.tables):
                table_data = table.extract()
                if table_data: # Убедимся, что из таблицы извлеклись данные
                    md_table = _convert_extracted_table_to_markdown(table_data)
                    if md_table: # Убедимся, что Markdown не пустой
                        rendered_tables[i] = md_table
                        table_insertion_points[i] = fitz.Rect(table.bbox).y0 # y0 - верхняя координата

            # Сортируем индексы таблиц по их вертикальному положению
            sorted_table_indices = sorted(table_insertion_points, key=table_insertion_points.get)
            yielded_tables = {idx: False for idx in sorted_table_indices} # Отслеживаем вставленные таблицы

            # 3. Получить текстовые блоки
            # Используем "blocks" для получения текста с координатами
            # sort=True упорядочивает блоки по y, затем по x (порядок чтения)
            blocks = page.get_text("blocks", sort=True)

            last_y_pos = 0 # Для отслеживания позиции вставки таблиц

            # 4. Обработать блоки и вставить таблицы
            for block in blocks:
                block_bbox = fitz.Rect(block[:4]) # Координаты блока (x0, y0, x1, y1)
                block_text = block[4].strip()     # Текст блока
                block_top_y = block_bbox.y0       # Верхняя координата блока

                # Проверяем, нужно ли вставить таблицу ПЕРЕД этим блоком
                for table_idx in sorted_table_indices:
                    # Если таблица еще не вставлена и ее верхняя граница выше или на уровне текущего блока
                    if not yielded_tables[table_idx] and table_insertion_points[table_idx] <= block_top_y:
                         # Дополнительная проверка, чтобы не вставлять таблицы слишком близко друг к другу
                         # или перед блоком, который является частью предыдущей таблицы.
                         # Порог (-5) можно настроить.
                         if table_insertion_points[table_idx] >= last_y_pos - 5:
                            page_elements.append(rendered_tables[table_idx])
                            yielded_tables[table_idx] = True
                            # Обновляем позицию, чтобы следующая таблица/текст вставлялись после этой таблицы
                            last_y_pos = max(last_y_pos, fitz.Rect(tabs.tables[table_idx].bbox).y1)


                # 5. Добавить текст блока, если он не является частью таблицы
                if block_text and not _is_block_inside_bbox(block_bbox, table_bboxes):
                    page_elements.append(block_text)
                    # Обновляем позицию последним обработанным текстовым блоком
                    last_y_pos = max(last_y_pos, block_bbox.y1)

            # 6. Вставить оставшиеся таблицы (если они находятся в самом конце страницы)
            for table_idx in sorted_table_indices:
                if not yielded_tables[table_idx]:
                    page_elements.append(rendered_tables[table_idx])

            # Объединяем элементы страницы в одну строку Markdown
            final_markdown_content += "\n\n".join(page_elements) + "\n\n" # Двойной перенос строки между элементами

        doc.close()
        logging.info(f"Successfully extracted Markdown with explicit tables from PDF: {pdf_path}")
        # Финальная очистка от лишних пустых строк
        final_markdown_content = "\n".join(line for line in final_markdown_content.splitlines() if line.strip())
        return final_markdown_content

    except Exception as e:
        logging.error(f"Error processing PDF file {pdf_path} with table extraction: {e}")
        # Можно вернуть пустую строку или None, или перевыбросить исключение
        # raise e
        return "" # Возвращаем пустую строку при ошибке

# --- Обработка DOCX ---

def extract_markdown_from_docx(docx_path: str) -> str:
    """
    Извлекает текст из DOCX и форматирует его в markdown.
    Обрабатывает параграфы и таблицы.
    """

    markdown_elements = []
    try:
        doc = docx.Document(docx_path)
        for element in doc.element.body:
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(element, doc)
                if para.text.strip(): # Дообавляем только непустые параграфы
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    markdown_elements.append(f"{'#' * level} {para.text.strip()}")
                else:
                    markdown_elements.append(para.text.strip())
            elif isinstance(element, docx.oxml.table.CT_Tbl):
                table = docx.table.Table(element, doc)
                if table.rows:
                    markdown_elements.append(_convert_table_to_markdown(table))

        logging.info(f"Successfully extracted Markdown from DOCX: {docx_path}")
        return "\n\n".join(markdown_elements)
    except Exception as e:
        logging.error(f"Error processing DOCX file {docx_path}: {e}")
        raise

# --- Обработка DOC ---

def convert_doc_to_docx(doc_path: str, output_dir: str) -> Optional[str]:
    """
    Конвертирует DOC в DOCX с помощью LibreOffice/soffice.
    Возвращает путь к сконвертированному файлу или None в случае ошибки.
    """
    if not Path(doc_path).exists():
         logging.error(f"DOC file not found: {doc_path}")
         return None

    try:
        # Имя выходного файла будет таким же, но с расширением .docx
        output_filename = Path(doc_path).stem + ".docx"
        output_path = Path(output_dir) / output_filename

        # Команда для конвертации
        # Путь к soffice может отличаться в вашей системе
        # На Windows это может быть что-то вроде "C:\Program Files\LibreOffice\program\soffice.exe"
        # На Linux/macOS часто просто "soffice" или "/usr/bin/soffice"
        cmd = [
            "soffice", # или полный путь к исполняемому файлу soffice
            "--headless", # Запуск без GUI
            "--convert-to", "docx",
            "--outdir", str(output_dir),
            str(doc_path)
        ]
        logging.info(f"Attempting to convert DOC to DOCX: {' '.join(cmd)}")

        # Запуск процесса конвертации
        result = subprocess.run(cmd, capture_output=True, text=True, check=False) # check=False, чтобы обработать ошибки вручную

        if result.returncode == 0 and output_path.exists():
            logging.info(f"Successfully converted {doc_path} to {output_path}")
            return str(output_path)
        else:
            logging.error(f"Failed to convert {doc_path}.")
            logging.error(f"Return Code: {result.returncode}")
            logging.error(f"Stderr: {result.stderr}")
            logging.error(f"Stdout: {result.stdout}")
            # Попытка удалить потенциально поврежденный выходной файл
            if output_path.exists():
                try:
                    os.remove(output_path)
                except OSError as rm_err:
                    logging.warning(f"Could not remove potentially incomplete output file {output_path}: {rm_err}")
            return None

    except FileNotFoundError:
        logging.error("`soffice` command not found. Is LibreOffice installed and in the system's PATH?")
        return None
    except Exception as e:
        logging.error(f"Error during DOC to DOCX conversion for {doc_path}: {e}")
        return None

def extract_markdown_from_doc(doc_path: str) -> Optional[str]:
    """
    Обрабатывает DOC файл: конвертирует в DOCX, затем извлекает Markdown.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        converted_docx_path = convert_doc_to_docx(doc_path, temp_dir)
        if converted_docx_path:
            try:
                markdown_text = extract_markdown_from_docx(converted_docx_path)
                # Опционально: удалить временный docx файл можно и не делать,
                # так как вся папка temp_dir будет удалена автоматически.
                return markdown_text
            except Exception as e:
                 logging.error(f"Failed to extract markdown from converted DOCX ({converted_docx_path}): {e}")
                 return None # Возвращаем None при ошибке на этапе извлечения из DOCX
        else:
            logging.error(f"Conversion failed for {doc_path}. Cannot extract markdown.")
            return None # Возвращаем None, если конвертация не удалась


# --- Главная функция-диспетчер ---

def process_document(file_path: str) -> Optional[str]:
    """
    Принимает путь к файлу (PDF, DOCX, DOC), извлекает текст
    и возвращает его в формате Markdown.
    Возвращает None в случае ошибки или неподдерживаемого формата.
    """
    path = Path(file_path)
    if not path.is_file():
        logging.error(f"File not found: {file_path}")
        return None

    file_ext = path.suffix.lower()
    markdown_content = None

    try:
        if file_ext == ".pdf":
            markdown_content = extract_markdown_from_pdf_with_tables(str(path))
        elif file_ext == ".docx":
            markdown_content = extract_markdown_from_docx(str(path))
        elif file_ext == ".doc":
            markdown_content = extract_markdown_from_doc(str(path))
        else:
            logging.warning(f"Unsupported file format: {file_ext} for file {file_path}")
            return None

        if markdown_content is None:
             logging.warning(f"Extraction resulted in None for file: {file_path}")

        return markdown_content

    except Exception as e:
        logging.error(f"An unexpected error occurred while processing {file_path}: {e}")
        return None # Возвращаем None при любой необработанной ошибке

#file_path = "D:\\BOT\\data\\contexts\\866070767\\test\\documents\\test_pdf.pdf"
#extracted_text = process_document(file_path)
#
#print("====================# Извлеченный текст #====================")
#print(extracted_text)
#print("=============================================================")
#
#txt_file_path = "D:\\BOT\\data\\contexts\\866070767\\test\\text\\test_pdf.txt"
#with open(txt_file_path, 'w', encoding="utf-8") as txt_out_file:
#    txt_out_file.write(extracted_text)
